# # agents.py
# # Multi-agent LangGraph pipeline for per-chunk legal review (dense-only RAG, Pinecone v2/v3 compatible).
# # Outputs: top-3 refs per chunk, structured issues, reviewed DOCX (with a Review Notes section), and a JSON report.

# import os
# import io
# import json
# from dataclasses import dataclass, field
# from typing import Any, Dict, List, Tuple, Optional

# # LangGraph (pip install langgraph langchain-core)
# from langgraph.graph import StateGraph, END
# from langgraph.checkpoint.memory import MemorySaver

# # LLM (OpenAI SDK >=1.0)
# import google.generativeai as genai
# _HAS_GEMINI = True


# # Embeddings / Reranker
# from sentence_transformers import SentenceTransformer, CrossEncoder
# import numpy as np

# # Pinecone v2/v3 compatibility
# PC_V3 = False
# try:
#     from pinecone import Pinecone, __version__ as _pcv
#     PC_V3 = True
# except Exception:
#     import pinecone  # v2
#     _pcv = getattr(pinecone, "__version__", "2.x")

# # DOCX writer
# from docx import Document
# from docx.shared import Pt
# from docx.enum.text import WD_COLOR_INDEX


# # ---------------- Tunables ----------------
# DENSE_MODEL = "BAAI/bge-base-en-v1.5"
# RERANKER_MODEL = "cross-encoder/ms-marco-MiniLM-L-6-v2"   # fast CPU
# TOP_K = 5            # first pass candidates per chunk
# TOP_N = 3            # final references per chunk
# BATCH_EMB = 32
# RERANK_BATCH = 32
# TEXT_CAP = 1600      # cap chunk text in prompts
# # ------------------------------------------


# # ---------------- State -------------------
# @dataclass
# class ChunkInput:
#     chunk_id: str
#     text: str
#     meta: Dict[str, Any]  # filename, section titles, etc.

# @dataclass
# class RefHit:
#     doc_id: Optional[str]
#     chunk_id: Optional[str]
#     title: Optional[str]
#     filename: Optional[str]
#     drive_link: Optional[str]
#     page_start: Optional[int]
#     page_end: Optional[int]
#     text: str
#     score: float

# @dataclass
# class AnalystIssue:
#     issue: str
#     severity: str
#     citation: str  # freeform citation string (we'll include drive link/pages)
#     suggestion: str

# @dataclass
# class ChunkResult:
#     chunk_id: str
#     top_refs: List[RefHit] = field(default_factory=list)
#     issues: List[AnalystIssue] = field(default_factory=list)
#     improved_text: Optional[str] = None  # writer agent suggestion

# @dataclass
# class PipelineState:
#     # constant across run
#     file_name: str
#     llm_model: str
#     pinecone_index_name: str
#     pinecone_namespace: Optional[str]
#     openai_api_key: Optional[str]

#     # runtime / shared clients
#     dense_model: Any = None
#     reranker: Any = None
#     pc_index: Any = None
#     llm_client: Any = None

#     # per-step / per-chunk
#     current_chunk: Optional[ChunkInput] = None
#     retrieved: List[RefHit] = field(default_factory=list)
#     analyst_issues: List[AnalystIssue] = field(default_factory=list)
#     writer_text: Optional[str] = None

#     # global outputs
#     results: List[ChunkResult] = field(default_factory=list)
# # ------------------------------------------


# # ------------- Helpers: init --------------
# def _init_clients(state: PipelineState):
#     # Dense encoder
#     if state.dense_model is None:
#         state.dense_model = SentenceTransformer(DENSE_MODEL)
#         state.dense_model.max_seq_length = 512
#     # Reranker
#     if state.reranker is None:
#         state.reranker = CrossEncoder(RERANKER_MODEL)
#     # Pinecone
#     api_key = os.getenv("PINECONE_API_KEY")
#     if PC_V3:
#         pc = Pinecone(api_key=api_key)
#         state.pc_index = pc.Index(state.pinecone_index_name)
#     else:
#         pinecone.init(api_key=api_key)
#         state.pc_index = pinecone.Index(state.pinecone_index_name)
#     # LLM
#     # LLM (Gemini)
#     if state.openai_api_key:  # carrying GOOGLE_API_KEY
#         genai.configure(api_key=state.openai_api_key)
#     # no client object needed; we’ll construct a GenerativeModel per call



# def _pick_text(md: Dict[str, Any]) -> str:
#     for k in ("text", "context", "content", "chunk", "body"):
#         v = md.get(k)
#         if isinstance(v, str) and v.strip():
#             return v
#     return ""


# def _encode_dense(texts: List[str], model: Any) -> List[List[float]]:
#     vecs = model.encode(texts, batch_size=BATCH_EMB, normalize_embeddings=True)
#     return [v.astype(np.float32).tolist() for v in vecs]


# def _pinecone_query_dense(index, vector: List[float], top_k: int, namespace: Optional[str]) -> List[Dict[str, Any]]:
#     q: Dict[str, Any] = dict(top_k=top_k, include_metadata=True)
#     if namespace:
#         q["namespace"] = namespace
#     if PC_V3:
#         res = index.query(vector=vector, **q)
#         matches = getattr(res, "matches", None) or res.get("matches", [])
#     else:
#         res = index.query(vector=vector, **q)
#         matches = res.get("matches", [])
#     out = []
#     for m in matches:
#         mid = getattr(m, "id", None) or m.get("id")
#         mscore = getattr(m, "score", None) or m.get("score")
#         mmeta = getattr(m, "metadata", None) or m.get("metadata", {})
#         out.append({"id": mid, "score": float(mscore), "metadata": mmeta})
#     return out


# def _rerank_pairs(reranker: Any, query_text: str, cands: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
#     pairs, keep = [], []
#     for it in cands:
#         txt = _pick_text(it["metadata"])
#         if txt:
#             pairs.append((query_text, txt))
#             keep.append(it)
#     if not pairs:
#         return []
#     logits = reranker.predict(pairs, batch_size=RERANK_BATCH, show_progress_bar=False)
#     sig = 1 / (1 + np.exp(-np.asarray(logits)))
#     ranked = []
#     for it, logit, s in zip(keep, logits, sig):
#         it = dict(it)
#         it["reranker_logit"] = float(logit)
#         it["reranker_score"] = float(s)
#         ranked.append(it)
#     ranked.sort(key=lambda x: x["reranker_score"], reverse=True)
#     return ranked


# def _map_hits(query_text: str, hits: List[Dict[str, Any]], top_n=TOP_N) -> List[RefHit]:
#     out: List[RefHit] = []
#     for h in hits[:top_n]:
#         md = h["metadata"] or {}
#         out.append(
#             RefHit(
#                 doc_id=md.get("doc_id"),
#                 chunk_id=md.get("chunk_id") or md.get("chunk_index"),
#                 title=md.get("doc_title") or md.get("title"),
#                 filename=md.get("file_name") or md.get("filename"),
#                 drive_link=md.get("drive_link"),
#                 page_start=md.get("page_start"),
#                 page_end=md.get("page_end"),
#                 text=_pick_text(md),
#                 score=float(h.get("reranker_score", h.get("score", 0.0))),
#             )
#         )
#     return out


# def _cap_text(s: str, n: int = TEXT_CAP) -> str:
#     s = (s or "").strip()
#     if len(s) <= n:
#         return s
#     return s[:n]


# # ---------- Nodes (LangGraph) -------------
# def node_retrieve(state: PipelineState) -> PipelineState:
#     _init_clients(state)
#     ch = state.current_chunk
#     assert ch is not None, "Missing current chunk"
#     qtext = _cap_text(ch.text)

#     dv = _encode_dense([qtext], state.dense_model)[0]
#     hits = _pinecone_query_dense(state.pc_index, dv, TOP_K, state.pinecone_namespace)
#     # rerank to pick best 3
#     ranked = _rerank_pairs(state.reranker, qtext, hits)
#     best = _map_hits(qtext, ranked, TOP_N)
#     state.retrieved = best
#     return state


# ANALYST_SYS = (
#     "You are a strict legal analyst for ADGM (Abu Dhabi Global Market). "
#     "Given a document chunk and up to 3 reference snippets from an ADGM-aligned knowledge base, "
#     "find legal issues/red flags and propose compliant fixes. "
#     "Output ONLY JSON with fields: issues=[{issue,severity,citation,suggestion}]. "
#     "Severity ∈ {Low,Medium,High}. Include precise citations using whatever metadata is present "
#     "(e.g., document title, drive link, pages)."
# )

# REVIEWER_SYS = (
#     "You are a meticulous legal reviewer. You're given the analyst's JSON and the same context. "
#     "Correct mistakes, remove weak flags, merge duplicates, and improve suggestions to be compliant with ADGM. "
#     "Return ONLY JSON with the same schema: issues=[{issue,severity,citation,suggestion}]."
# )

# WRITER_SYS = (
#     "You are a careful editor. Given the original chunk and the reviewed issues, propose improved text "
#     "that keeps the original meaning but makes it compliant. Output ONLY JSON: {improved_text: string}."
# )

# def _llm_json(_unused_client: Any, model: str, system: str, user: str) -> Dict[str, Any]:
#     # Build a Gemini model with the role’s system instruction
#     gmodel = genai.GenerativeModel(
#         model_name=model,
#         system_instruction=system,
#     )
#     resp = gmodel.generate_content(
#         [user],
#         generation_config={
#             "temperature": 0.2,
#             "response_mime_type": "application/json",
#         },
#     )
#     text = (resp.text or "").strip()
#     try:
#         return json.loads(text)
#     except Exception:
#         start, end = text.find("{"), text.rfind("}")
#         if start != -1 and end != -1 and end > start:
#             try:
#                 return json.loads(text[start : end + 1])
#             except Exception:
#                 pass
#         # robust fallback by role
#         return {"issues": []} if "issues" in system.lower() else {"improved_text": ""}



# def node_analyst(state: PipelineState) -> PipelineState:
#     ch = state.current_chunk
#     refs = state.retrieved
#     # Build user prompt
#     bundle = {
#         "chunk": _cap_text(ch.text),
#         "references": [
#             {
#                 "title": r.title,
#                 "filename": r.filename,
#                 "drive_link": r.drive_link,
#                 "pages": [r.page_start, r.page_end],
#                 "text": _cap_text(r.text, 1200),
#             }
#             for r in refs
#         ],
#         # Known ADGM red-flags guidance (anchors model)
#         "known_red_flags": [
#             "Jurisdiction not ADGM or courts unspecified",
#             "Missing/invalid signatory or witness section",
#             "Missing core incorporation info (share capital, registered office, directors)",
#             "Ambiguous or non-binding language in obligations",
#             "Templates not aligned with ADGM Companies Regulations 2020",
#         ],
#     }
#     user = json.dumps(bundle, ensure_ascii=False)
#     out = _llm_json(state.llm_client, state.llm_model, ANALYST_SYS, user)
#     issues = []
#     for it in out.get("issues", []):
#         issues.append(
#             AnalystIssue(
#                 issue=str(it.get("issue", "")).strip(),
#                 severity=str(it.get("severity", "Medium")).strip() or "Medium",
#                 citation=str(it.get("citation", "")).strip(),
#                 suggestion=str(it.get("suggestion", "")).strip(),
#             )
#         )
#     state.analyst_issues = issues
#     return state


# def node_reviewer(state: PipelineState) -> PipelineState:
#     ch = state.current_chunk
#     refs = state.retrieved
#     payload = {
#         "chunk": _cap_text(ch.text),
#         "references": [
#             {
#                 "title": r.title,
#                 "filename": r.filename,
#                 "drive_link": r.drive_link,
#                 "pages": [r.page_start, r.page_end],
#                 "text": _cap_text(r.text, 1200),
#             }
#             for r in refs
#         ],
#         "analyst": [{"issue": i.issue, "severity": i.severity, "citation": i.citation, "suggestion": i.suggestion} for i in state.analyst_issues],
#     }
#     user = json.dumps(payload, ensure_ascii=False)
#     out = _llm_json(state.llm_client, state.llm_model, REVIEWER_SYS, user)
#     # If reviewer returns nothing, keep analyst
#     issues = []
#     arr = out.get("issues", []) if isinstance(out, dict) else []
#     if not arr:
#         arr = [{"issue": i.issue, "severity": i.severity, "citation": i.citation, "suggestion": i.suggestion} for i in state.analyst_issues]
#     for it in arr:
#         issues.append(
#             AnalystIssue(
#                 issue=str(it.get("issue", "")).strip(),
#                 severity=str(it.get("severity", "Medium")).strip() or "Medium",
#                 citation=str(it.get("citation", "")).strip(),
#                 suggestion=str(it.get("suggestion", "")).strip(),
#             )
#         )
#     state.analyst_issues = issues
#     return state


# def node_writer(state: PipelineState) -> PipelineState:
#     ch = state.current_chunk
#     payload = {
#         "chunk": _cap_text(ch.text, 1600),
#         "issues": [{"issue": i.issue, "severity": i.severity, "suggestion": i.suggestion} for i in state.analyst_issues],
#     }
#     user = json.dumps(payload, ensure_ascii=False)
#     out = _llm_json(state.llm_client, state.llm_model, WRITER_SYS, user)
#     state.writer_text = str(out.get("improved_text", "")).strip()
#     # Persist chunk result
#     state.results.append(
#         ChunkResult(
#             chunk_id=ch.chunk_id,
#             top_refs=state.retrieved,
#             issues=state.analyst_issues,
#             improved_text=state.writer_text or None,
#         )
#     )
#     # clear per-chunk working buffers
#     state.retrieved = []
#     state.analyst_issues = []
#     state.writer_text = None
#     state.current_chunk = None
#     return state
# # ------------------------------------------


# # ---------- Graph assembly ---------------
# def _build_graph():
#     g = StateGraph(PipelineState)
#     g.add_node("retrieve", node_retrieve)
#     g.add_node("analyst", node_analyst)
#     g.add_node("reviewer", node_reviewer)
#     g.add_node("writer", node_writer)

#     g.set_entry_point("retrieve")
#     g.add_edge("retrieve", "analyst")
#     g.add_edge("analyst", "reviewer")
#     g.add_edge("reviewer", "writer")
#     g.add_edge("writer", END)

#     # memory = MemorySaver()
#     return g.compile()


# GRAPH = _build_graph()
# # ------------------------------------------


# # ---------- Orchestrators -----------------
# def run_for_chunks(
#     file_name: str,
#     chunks: List[Dict[str, Any]],
#     pinecone_index_name: str,
#     pinecone_namespace: Optional[str],
#     llm_model: str = None,
#     openai_api_key: Optional[str] = None,
# ) -> Dict[str, Any]:
#     """Runs graph per chunk, returns {'results': List[ChunkResult], 'report': dict}."""
#     llm_model = llm_model or os.getenv("OPENAI_MODEL", "gpt-4o-mini")

#     base_state = PipelineState(
#         file_name=file_name,
#         llm_model=llm_model,
#         pinecone_index_name=pinecone_index_name,
#         pinecone_namespace=pinecone_namespace,
#         openai_api_key=openai_api_key or os.getenv("OPENAI_API_KEY"),
#     )
#     _init_clients(base_state)

#     # Feed each chunk through the graph
#     for c in chunks:
#         ch = ChunkInput(
#             chunk_id=c.get("id"),
#             text=c.get("text") or "",
#             meta=c.get("metadata") or {},
#         )
#         s = PipelineState(**{**base_state.__dict__})
#         s.current_chunk = ch
#     #     cfg = {"configurable": {
#     #     "thread_id": f"{file_name}:{c.get('id')}",   # unique per chunk
#     #     "checkpoint_ns": "per-chunk"                 # optional but nice to have
#     # }}
#         out_state = GRAPH.invoke(s)
#         # Append accumulated result already handled in writer node

#     # Build a summary report
#     issues_flat = []
#     for r in base_state.results:
#         for it in r.issues:
#             # compute a basic citation string from refs if missing
#             citation = it.citation
#             if not citation and r.top_refs:
#                 r0 = r.top_refs[0]
#                 citation = f"{r0.title or r0.filename or ''} {f'({r0.page_start}-{r0.page_end})' if r0.page_start else ''} {r0.drive_link or ''}".strip()
#             issues_flat.append(
#                 {
#                     "chunk_id": r.chunk_id,
#                     "issue": it.issue,
#                     "severity": it.severity,
#                     "citation": citation,
#                     "suggestion": it.suggestion,
#                 }
#             )

#     report = {
#         "file": file_name,
#         "chunks_processed": len(chunks),
#         "issues_found": len(issues_flat),
#         "items": issues_flat,
#     }
#     return {"results": base_state.results, "report": report}


# def build_reviewed_docx(
#     original_filename: str,
#     chunks: List[Dict[str, Any]],
#     results: List[ChunkResult],
# ) -> bytes:
#     """
#     Create a new DOCX with a 'Review Notes' section containing:
#     - Chunk preview
#     - Issues with severity + citation
#     - Writer agent improved text (if present)
#     This avoids fragile true Word comments and is portable.
#     """
#     doc = Document()
#     doc.add_heading(f"AI Review Notes for {original_filename}", level=1)

#     # Simple map for quick lookup
#     rmap: Dict[str, ChunkResult] = {r.chunk_id: r for r in results}

#     for c in chunks:
#         rid = c.get("id")
#         res = rmap.get(rid)
#         if not res or (not res.issues and not res.improved_text):
#             continue

#         # Section header
#         h = doc.add_heading("", level=2)
#         run = h.add_run(f"Section / Chunk {rid}")
#         run.font.size = Pt(12)

#         # Preview
#         p = doc.add_paragraph()
#         run = p.add_run("Preview: ")
#         run.bold = True
#         p.add_run((c.get("preview") or (c.get("text") or "")[:160]).strip())

#         # Issues
#         if res.issues:
#             doc.add_paragraph("Issues:", style=None)
#             for idx, it in enumerate(res.issues, 1):
#                 par = doc.add_paragraph(f"{idx}. [{it.severity}] {it.issue}")
#                 if it.citation:
#                     citation = doc.add_paragraph()
#                     cit_run = citation.add_run(f"  Citation: {it.citation}")
#                     cit_run.font.size = Pt(9)
#                 if it.suggestion:
#                     s = doc.add_paragraph()
#                     s_run = s.add_run(f"  Suggestion: {it.suggestion}")
#                     s_run.font.size = Pt(10)
#                     s_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

#         # Improved text
#         if res.improved_text:
#             doc.add_paragraph("Proposed Compliant Text:", style=None)
#             t = doc.add_paragraph(res.improved_text)

#         # Divider
#         doc.add_paragraph("-" * 40)

#     bio = io.BytesIO()
#     doc.save(bio)
#     bio.seek(0)
#     return bio.read()
# # ------------------------------------------



# agents.py
# Multi-agent pipeline (LangGraph) for per-chunk legal review.
# Dense-only Pinecone retrieval (v2 or v3 client), rerank with MiniLM,
# LLM provider: auto (Gemini if GOOGLE_API_KEY, else OpenAI).
# Writer adds a "Review Notes" section with issues, citations, and verbatim excerpts.

import os
import io
import json
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional

import numpy as np
from sentence_transformers import SentenceTransformer, CrossEncoder

# Pinecone v2/v3 compatibility
PC_V3 = False
try:
    from pinecone import Pinecone, __version__ as _pcv
    PC_V3 = True
except Exception:
    import pinecone  # v2
    _pcv = getattr(pinecone, "__version__", "2.x")

# LangGraph (no checkpointer to avoid msgpack serialization issues)
from langgraph.graph import StateGraph, END

# LLM auto-select
USE_GEMINI = bool(os.getenv("GOOGLE_API_KEY"))
if USE_GEMINI:
    import google.generativeai as genai
else:
    from openai import OpenAI

# DOCX writer
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX


# ---------------- Tunables ----------------
DENSE_MODEL = "BAAI/bge-base-en-v1.5"
RERANKER_MODEL = "cross-encoder/ms-marco-MiniLM-L-6-v2"   # fast CPU
TOP_K = 5            # candidates per chunk (dense-only)
TOP_N = 3            # top references sent to LLM and used for citations
BATCH_EMB = 32
RERANK_BATCH = 32
TEXT_CAP = 1600      # cap chunk/ref text in prompts & docx
EXCERPT_LEN = 320    # verbatim excerpt length for DOCX
# ------------------------------------------


# ---------------- State -------------------
@dataclass
class ChunkInput:
    chunk_id: str
    text: str
    meta: Dict[str, Any]

@dataclass
class RefHit:
    doc_id: Optional[str]
    chunk_id: Optional[str]
    title: Optional[str]
    filename: Optional[str]
    drive_link: Optional[str]
    page_start: Optional[int]
    page_end: Optional[int]
    text: str
    score: float

@dataclass
class AnalystIssue:
    issue: str
    severity: str
    citation: str
    suggestion: str
    supporting_refs: List[int] = field(default_factory=list)  # indices into provided references

@dataclass
class ChunkResult:
    chunk_id: str
    top_refs: List[RefHit] = field(default_factory=list)
    issues: List[AnalystIssue] = field(default_factory=list)
    improved_text: Optional[str] = None  # writer agent suggestion

@dataclass
class PipelineState:
    # config
    file_name: str
    llm_model: str
    pinecone_index_name: str
    pinecone_namespace: Optional[str]
    api_key: Optional[str]  # GOOGLE_API_KEY or OPENAI_API_KEY depending on provider

    # runtime singletons (kept simple; no checkpointer)
    dense_model: Any = None
    reranker: Any = None
    pc_index: Any = None
    llm_client: Any = None  # OpenAI client or None for Gemini

    # per-chunk scratch
    current_chunk: Optional[ChunkInput] = None
    retrieved: List[RefHit] = field(default_factory=list)
    analyst_issues: List[AnalystIssue] = field(default_factory=list)
    writer_text: Optional[str] = None

    # outputs
    results: List[ChunkResult] = field(default_factory=list)
# ------------------------------------------


# ---------- Init / helpers ----------
def _init_clients(state: PipelineState):
    # Dense encoder
    if state.dense_model is None:
        state.dense_model = SentenceTransformer(DENSE_MODEL)
        state.dense_model.max_seq_length = 512
    # Reranker
    if state.reranker is None:
        state.reranker = CrossEncoder(RERANKER_MODEL)
    # Pinecone
    if PC_V3:
        pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))
        state.pc_index = pc.Index(state.pinecone_index_name)
    else:
        pinecone.init(api_key=os.getenv("PINECONE_API_KEY"))
        state.pc_index = pinecone.Index(state.pinecone_index_name)
    # LLM
    if USE_GEMINI:
        genai.configure(api_key=state.api_key)
        state.llm_client = None  # not needed for Gemini
    else:
        state.llm_client = OpenAI(api_key=state.api_key)


def _pick_text(md: Dict[str, Any]) -> str:
    for k in ("text", "context", "content", "chunk", "body"):
        v = md.get(k)
        if isinstance(v, str) and v.strip():
            return v
    return ""


def _encode_dense(texts: List[str], model: Any) -> List[List[float]]:
    vecs = model.encode(texts, batch_size=BATCH_EMB, normalize_embeddings=True)
    return [v.astype(np.float32).tolist() for v in vecs]


def _pinecone_query_dense(index, vector: List[float], top_k: int, namespace: Optional[str]) -> List[Dict[str, Any]]:
    q: Dict[str, Any] = dict(top_k=top_k, include_metadata=True)
    if namespace:
        q["namespace"] = namespace
    if PC_V3:
        res = index.query(vector=vector, **q)
        matches = getattr(res, "matches", None) or res.get("matches", [])
    else:
        res = index.query(vector=vector, **q)
        matches = res.get("matches", [])
    out = []
    for m in matches:
        mid = getattr(m, "id", None) or m.get("id")
        mscore = getattr(m, "score", None) or m.get("score")
        mmeta = getattr(m, "metadata", None) or m.get("metadata", {})
        out.append({"id": mid, "score": float(mscore), "metadata": mmeta})
    return out


def _rerank_pairs(reranker: Any, query_text: str, cands: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    pairs, keep = [], []
    for it in cands:
        txt = _pick_text(it["metadata"])
        if txt:
            pairs.append((query_text, txt))
            keep.append(it)
    if not pairs:
        return []
    logits = reranker.predict(pairs, batch_size=RERANK_BATCH, show_progress_bar=False)
    sig = 1 / (1 + np.exp(-np.asarray(logits)))
    ranked = []
    for it, logit, s in zip(keep, logits, sig):
        it = dict(it)
        it["reranker_logit"] = float(logit)
        it["reranker_score"] = float(s)
        ranked.append(it)
    ranked.sort(key=lambda x: x["reranker_score"], reverse=True)
    return ranked


def _map_hits(hits: List[Dict[str, Any]], top_n=TOP_N) -> List[RefHit]:
    out: List[RefHit] = []
    for h in hits[:top_n]:
        md = h["metadata"] or {}
        out.append(
            RefHit(
                doc_id=md.get("doc_id"),
                chunk_id=md.get("chunk_id") or md.get("chunk_index"),
                title=md.get("doc_title") or md.get("title"),
                filename=md.get("file_name") or md.get("filename"),
                drive_link=md.get("drive_link"),
                page_start=md.get("page_start"),
                page_end=md.get("page_end"),
                text=_pick_text(md),
                score=float(h.get("reranker_score", h.get("score", 0.0))),
            )
        )
    return out


def _cap(s: str, n: int = TEXT_CAP) -> str:
    s = (s or "").strip()
    return s if len(s) <= n else s[:n]


# ---------- Prompts ----------
ANALYST_SYS = (
    "You are a strict legal analyst for ADGM. Given a document chunk and up to 3 reference snippets from an "
    "ADGM-aligned knowledge base, identify legal issues/red flags and propose compliant fixes. "
    "Return ONLY JSON:\n"
    "{ \"issues\": [ { \"issue\": str, \"severity\": \"Low|Medium|High\", \"citation\": str, "
    "\"suggestion\": str, \"supporting_refs\": [int] } ] }\n"
    "Notes:\n"
    "- 'supporting_refs' are ZERO-based indices into the provided 'references' array that support the issue.\n"
    "- 'citation' should be a readable string that mentions the document title and page range if known.\n"
)

REVIEWER_SYS = (
    "You are a meticulous legal reviewer. You get the original chunk, the references, and the analyst's JSON. "
    "Correct mistakes, remove weak flags, merge duplicates, and improve suggestions to be ADGM-compliant. "
    "Return ONLY JSON with the same schema as the analyst."
)

WRITER_SYS = (
    "You are a careful editor. Given the original chunk and the reviewed issues, propose improved text that "
    "keeps the original meaning but fixes problems. Output ONLY JSON: { \"improved_text\": str }"
)


# ---------- LLM wrapper ----------
def _llm_json(model_name: str, system: str, user: str, state: PipelineState) -> Dict[str, Any]:
    if USE_GEMINI:
        gmodel = genai.GenerativeModel(model_name, system_instruction=system)
        resp = gmodel.generate_content(
            [user],
            generation_config={"temperature": 0.2, "response_mime_type": "application/json"},
        )
        text = (resp.text or "").strip()
    else:
        resp = state.llm_client.chat.completions.create(
            model=model_name,
            messages=[{"role": "system", "content": system}, {"role": "user", "content": user}],
            temperature=0.2,
        )
        text = resp.choices[0].message.content.strip()

    try:
        return json.loads(text)
    except Exception:
        start, end = text.find("{"), text.rfind("}")
        if start != -1 and end != -1 and end > start:
            try:
                return json.loads(text[start : end + 1])
            except Exception:
                pass
        # role-based fallback
        return {"issues": []} if "issues" in system.lower() else {"improved_text": ""}


# ---------- Graph nodes ----------
def node_retrieve(state: PipelineState) -> PipelineState:
    _init_clients(state)
    ch = state.current_chunk
    assert ch is not None
    qtext = _cap(ch.text)
    dv = _encode_dense([qtext], state.dense_model)[0]
    hits = _pinecone_query_dense(state.pc_index, dv, TOP_K, state.pinecone_namespace)
    ranked = _rerank_pairs(state.reranker, qtext, hits)
    state.retrieved = _map_hits(ranked, TOP_N)
    return state


def node_analyst(state: PipelineState) -> PipelineState:
    ch = state.current_chunk
    refs = state.retrieved
    payload = {
        "chunk": _cap(ch.text),
        "references": [
            {
                "title": r.title,
                "filename": r.filename,
                "drive_link": r.drive_link,
                "pages": [r.page_start, r.page_end],
                "text": _cap(r.text, TEXT_CAP),
            }
            for r in refs
        ],
        "known_red_flags": [
            "Jurisdiction not ADGM or courts unspecified",
            "Missing/invalid signatory or witness section",
            "Missing core incorporation info (share capital, registered office, directors)",
            "Ambiguous language in obligations",
            "Not aligned with ADGM Companies Regulations 2020",
        ],
    }
    out = _llm_json(state.llm_model, ANALYST_SYS, json.dumps(payload), state)
    issues: List[AnalystIssue] = []
    for it in out.get("issues", []):
        refs_idx = it.get("supporting_refs") or []
        if isinstance(refs_idx, list):
            refs_idx = [int(x) for x in refs_idx if isinstance(x, (int, float))]
        issues.append(
            AnalystIssue(
                issue=str(it.get("issue", "")).strip(),
                severity=(str(it.get("severity", "Medium")).strip() or "Medium"),
                citation=str(it.get("citation", "")).strip(),
                suggestion=str(it.get("suggestion", "")).strip(),
                supporting_refs=refs_idx[:3],
            )
        )
    state.analyst_issues = issues
    return state


def node_reviewer(state: PipelineState) -> PipelineState:
    ch = state.current_chunk
    refs = state.retrieved
    payload = {
        "chunk": _cap(ch.text),
        "references": [
            {
                "title": r.title,
                "filename": r.filename,
                "drive_link": r.drive_link,
                "pages": [r.page_start, r.page_end],
                "text": _cap(r.text, TEXT_CAP),
            }
            for r in refs
        ],
        "analyst": [
            {
                "issue": i.issue,
                "severity": i.severity,
                "citation": i.citation,
                "suggestion": i.suggestion,
                "supporting_refs": i.supporting_refs,
            }
            for i in state.analyst_issues
        ],
    }
    out = _llm_json(state.llm_model, REVIEWER_SYS, json.dumps(payload), state)
    arr = out.get("issues", []) if isinstance(out, dict) else []
    if not arr:
        arr = [
            {"issue": i.issue, "severity": i.severity, "citation": i.citation, "suggestion": i.suggestion, "supporting_refs": i.supporting_refs}
            for i in state.analyst_issues
        ]
    issues: List[AnalystIssue] = []
    for it in arr:
        refs_idx = it.get("supporting_refs") or []
        if isinstance(refs_idx, list):
            refs_idx = [int(x) for x in refs_idx if isinstance(x, (int, float))]
        issues.append(
            AnalystIssue(
                issue=str(it.get("issue", "")).strip(),
                severity=(str(it.get("severity", "Medium")).strip() or "Medium"),
                citation=str(it.get("citation", "")).strip(),
                suggestion=str(it.get("suggestion", "")).strip(),
                supporting_refs=refs_idx[:3],
            )
        )
    state.analyst_issues = issues
    return state


def node_writer(state: PipelineState) -> PipelineState:
    ch = state.current_chunk
    payload = {
        "chunk": _cap(ch.text, TEXT_CAP),
        "issues": [{"issue": i.issue, "severity": i.severity, "suggestion": i.suggestion} for i in state.analyst_issues],
    }
    out = _llm_json(state.llm_model, WRITER_SYS, json.dumps(payload), state)
    state.writer_text = str(out.get("improved_text", "")).strip()
    state.results.append(
        ChunkResult(
            chunk_id=ch.chunk_id,
            top_refs=state.retrieved[:],
            issues=state.analyst_issues[:],
            improved_text=state.writer_text or None,
        )
    )
    state.retrieved = []
    state.analyst_issues = []
    state.writer_text = None
    state.current_chunk = None
    return state


# ---------- Graph assembly ----------
def _build_graph():
    g = StateGraph(PipelineState)
    g.add_node("retrieve", node_retrieve)
    g.add_node("analyst", node_analyst)
    g.add_node("reviewer", node_reviewer)
    g.add_node("writer", node_writer)

    g.set_entry_point("retrieve")
    g.add_edge("retrieve", "analyst")
    g.add_edge("analyst", "reviewer")
    g.add_edge("reviewer", "writer")
    g.add_edge("writer", END)

    return g.compile()  # NO checkpointer
GRAPH = _build_graph()


# ---------- Orchestrator ----------
def run_for_chunks(
    file_name: str,
    chunks: List[Dict[str, Any]],
    pinecone_index_name: str,
    pinecone_namespace: Optional[str],
    llm_model: str = None,
    openai_api_key: Optional[str] = None,  # reused for GOOGLE_API_KEY when using Gemini
) -> Dict[str, Any]:
    model = llm_model or (os.getenv("GEMINI_MODEL") if USE_GEMINI else os.getenv("OPENAI_MODEL", "gpt-4o-mini"))

    base = PipelineState(
        file_name=file_name,
        llm_model=model,
        pinecone_index_name=pinecone_index_name,
        pinecone_namespace=pinecone_namespace,
        api_key=openai_api_key or (os.getenv("GOOGLE_API_KEY") if USE_GEMINI else os.getenv("OPENAI_API_KEY")),
    )
    _init_clients(base)

    all_results: List[ChunkResult] = []
    for c in chunks:
        ch = ChunkInput(
            chunk_id=c.get("id"),
            text=c.get("text") or "",
            meta=c,
        )
        s = PipelineState(**{**base.__dict__})
        s.current_chunk = ch
        out_state = GRAPH.invoke(s)
        all_results.extend(out_state.results)

    # Build flat report
    items = []
    for r in all_results:
        for it in r.issues:
            # Build readable citation if missing
            citation = it.citation
            if not citation and r.top_refs:
                r0 = r.top_refs[0]
                pages = f" (pp. {r0.page_start}-{r0.page_end})" if r0.page_start else ""
                citation = f"{r0.title or r0.filename or 'KB Doc'}{pages} {r0.drive_link or ''}".strip()
            items.append(
                {
                    "chunk_id": r.chunk_id,
                    "issue": it.issue,
                    "severity": it.severity,
                    "suggestion": it.suggestion,
                    "citation": citation,
                    "supporting_refs": it.supporting_refs,
                }
            )

    report = {"file": file_name, "chunks_processed": len(chunks), "issues_found": len(items), "items": items}
    return {"results": all_results, "report": report}


# ---------- DOCX builder with verbatim references ----------
def build_reviewed_docx(
    original_filename: str,
    chunks: List[Dict[str, Any]],
    results: List[ChunkResult],
) -> bytes:
    doc = Document()
    doc.add_heading(f"AI Review Notes for {original_filename}", level=1)
    doc.add_paragraph("This section contains issues, suggested fixes, and verbatim excerpts with citations from the knowledge base.")

    # Quick map
    rmap: Dict[str, ChunkResult] = {r.chunk_id: r for r in results}

    for c in chunks:
        rid = c.get("id")
        res = rmap.get(rid)
        if not res or (not res.issues and not res.improved_text):
            continue

        # Section header
        h = doc.add_heading(f"Section / Chunk {rid}", level=2)
        h_run = h.runs[0]
        h_run.font.size = Pt(12)

        # Preview
        p = doc.add_paragraph()
        p.add_run("Preview: ").bold = True
        p.add_run((c.get("preview") or (c.get("text") or "")[:180]).strip())

        # References (top-3)
        if res.top_refs:
            doc.add_paragraph("References used:", style=None)
            for idx, ref in enumerate(res.top_refs):
                line = f"{idx}. {ref.title or ref.filename or 'KB Doc'}"
                if ref.page_start:
                    line += f" (pp. {ref.page_start}-{ref.page_end})"
                if ref.drive_link:
                    line += f" — {ref.drive_link}"
                doc.add_paragraph(line)

        # Issues with verbatim excerpts
        if res.issues:
            doc.add_paragraph("Issues & Suggested Fixes:", style=None)
            for k, it in enumerate(res.issues, 1):
                doc.add_paragraph(f"{k}. [{it.severity}] {it.issue}")
                if it.citation:
                    cit = doc.add_paragraph(f"   Citation: {it.citation}")
                    cit.runs[0].font.size = Pt(9)
                if it.suggestion:
                    sg = doc.add_paragraph(f"   Suggestion: {it.suggestion}")
                    sg.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
                    sg.runs[0].font.size = Pt(10)

                # Verbatim excerpts from supporting refs (fallback to ref 0 if not provided)
                support_idxs = it.supporting_refs or [0]
                doc.add_paragraph("   Evidence excerpts:", style=None)
                for si in support_idxs:
                    if 0 <= si < len(res.top_refs):
                        ref = res.top_refs[si]
                        excerpt = (_cap(ref.text, EXCERPT_LEN) or "").strip()
                        if excerpt:
                            q = doc.add_paragraph(f'   • "{excerpt}"')
                            q.runs[0].font.size = Pt(9)
                            meta = []
                            if ref.title or ref.filename:
                                meta.append(ref.title or ref.filename)
                            if ref.page_start:
                                meta.append(f"pp. {ref.page_start}-{ref.page_end}")
                            if ref.drive_link:
                                meta.append(ref.drive_link)
                            if meta:
                                meta_p = doc.add_paragraph("     — " + " · ".join(meta))
                                meta_p.runs[0].font.size = Pt(9)

        # Proposed text
        if res.improved_text:
            doc.add_paragraph("Proposed Compliant Text:", style=None)
            t = doc.add_paragraph(res.improved_text)

        # Divider
        doc.add_paragraph("-" * 60)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()
